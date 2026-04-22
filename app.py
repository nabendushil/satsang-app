from flask import Flask, render_template, request, redirect, send_file, jsonify
import os
import json
from datetime import datetime
from docx import Document

app = Flask(__name__)

# -------------------------------
# MASTER DATA (UNCHANGED)
# -------------------------------

namelist = [
"Nirmallya Shil","Narayan Ch. Shil","Gopal Ch. Sarkar","Haradhan Das","Sarajit Das",
"Sanjoy Roy","Tapan Ghosh","Dilip Roy","Rabi Roy",
"Rabindra Nath Sarkar","Sankar Malakar","Shibu Sharma","Jayanti Shil","Gopal Bala",
"Nityananda Sarkar","Prakash Sarkar","Prakash Sarkar",
"Alo Shil","Tapan Baidya","Parimal Dutta","Lotika Das","Gita Sarkar",
"Prasanta Sarkar","Sandhya Shil","Rajat Dey","Dulal Chanda","Bandana Sikdar",
"Ranjit Kr. Baidya","Bulu Dutta","Nimaichand Biswas","Bijay Biswas","Prasanta Roy",
"Gobinda Dutta","Sefali Dutta","Swapan Baidya","Utpal Talukdar","Kamal Baidya",
"Gopal Sarkar","Nikhil Basu","Amar Bala","Krishna Baidya","Bappa Biswakarma",
"Bipul Modak","Ram Goswami","Subrata Shil","Jyotsna Goswami","Sujoy Biswas",
"Tapash Sarkar","Sipra Saha","Manotosh Palit","Padmabati Mitra","Asha Pramanik",
"Paritosh Bhattacharjee","Sambhu Thakur","Subrata Sarkar","Gobinda Sarkar","Bela Das",
"Sachin Sarkar","Subrata Sarkar","Nripen Biswas","Anil Sarkar","Babu Sarkar",
"Nitya Gopal Bala","Gobinda Dutta","Swapan Baidya","Dipak Sarkar","Sunil Sarkar",
"Raj Goswami","Subrata Malakar","Sampa Biswas","Kalyan Sundar Dutta","Nirmal Mandal",
"Kishori Mandal","Ranjan Shil","Ashish Sanyal","Gouranga Karmakar","Mangal Karmakar",
"Bishnu Karmakar","Bipul Modak","Sadhan Samadder","Partha Mandal","Sudip Some (Lakhi)"
]

addresslist = [
"Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa",
"Jyotinagar, Phansidewa","Ramganj, Uttar Dinajpur","Jyotinagar, Phansidewa","Chathat, Phansidewa",
"Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa",
"Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Guabari, Phansidewa",
"Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa",
"Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa",
"Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa",
"Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa",
"Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa",
"Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa",
"Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa",
"Jyotinagar, Phansidewa","Bhaktinagar, Phansidewa","Bhaktinagar, Phansidewa","Bhaktinagar, Phansidewa",
"Bhaktinagar, Phansidewa","Bhaktinagar, Phansidewa","Bhaktinagar, Phansidewa","Bhaktinagar, Phansidewa",
"Bhaktinagar, Phansidewa","Bhaktinagar, Phansidewa","Chaturagachh, Jalpaiguri","Dundiajote, Phansidewa",
"Ambari, Ghoshpukur","Demdakhari, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa",
"Jyotinagar, Phansidewa","Demdakhari, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa",
"Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa",
"Jyotinagar, Phansidewa","Bhaktinagar, Phansidewa","Bhaktinagar, Phansidewa","Bhaktinagar, Phansidewa",
"Bhaktinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa","Jyotinagar, Phansidewa",
"Jyotinagar, Phansidewa","Bhaktinagar, Phansidewa","Bhaktinagar, Phansidewa","Bhaktinagar, Phansidewa",
"Bhaktinagar, Phansidewa","Bhaktinagar, Phansidewa","Bhaktinagar, Phansidewa","Palpakuria, Barasat",
"Ghospara, Rangapani"
]
  
familycode = [
"4417144700","5193400","1159917000","51497634810","51497798237",
"51497679992","833740300","51497569137","4403363000","51497616918",
"2631115100","1442222000","51497876104","51497862647","51497559268",
"51497979948","51107704200","51107704200","4402019800","4402019800",
"51497885611","2705436500","50018654600","51497879325","51497587338",
"51497773751","51498079746","51497789907","2027194800","51498004163",
"2098215100","50763372600","1001535300","51497699864","3551789100",
"3294047500","2733940800","50677945000","1264826400","50670923000",
"51497699874","51498018600","4412938600","51497890329","51497665900",
"3911691300","51497950398","50763370000","51498031375","51497950396",
"51497940741","51498082875","51497646090","51066935300","51498018610",
"51498005758","51497682984","51498037321","50629422700","51498064580",
"5705920900","51498037321","51497625356","967529900","51497970327",
"50845880400","003294047500","050677945000","510976248","2012638700",
"051497821504","51498064584","51498004161","51485949100","51498028830",
"51497876110","51497609214","51497774331","51497569133","51497664555",
"51497924187","003911691300","50845878200","51498086643","051498029526"
]


familycode = [str(code).zfill(12) for code in familycode]

# -------------------------------
# LOAD / SAVE
# -------------------------------

def load_data():
    try:
        with open("data.json", "r") as f:
            return json.load(f)
    except:
        return []

def save_data(data):
    with open("data.json", "w") as f:
        json.dump(data, f, indent=4)

def load_phones():
    try:
        with open("phones.json", "r") as f:
            return json.load(f)
    except:
        return {}

def save_phones(data):
    with open("phones.json", "w") as f:
        json.dump(data, f, indent=4)

# -------------------------------
# HOME
# -------------------------------

@app.route("/")
def home():
    return render_template("home.html")

# -------------------------------
# 📈 STATISTICS (FIX)
# -------------------------------

@app.route("/statistics")
def statistics():

    data = load_data()

    try:
        with open("month_status.json", "r") as f:
            monthly_status = json.load(f)
    except:
        monthly_status = {}

    return render_template(
        "statistics.html",
        data=data,
        monthly_status=monthly_status
    )
    
@app.route("/calendar")
def calendar():
    return render_template("calendar.html")

@app.route("/satsangi")
def satsangi():
    return render_template("satsangi.html", names=namelist)

# -------------------------------
# 🔍 SATSANGI SEARCH (SAFE FIX)
# -------------------------------

@app.route("/satsangi_search", methods=["GET", "POST"])
def satsangi_search():

    # 🛑 Prevent direct URL access (GET request)
    if request.method == "GET":
        return redirect("/satsangi")

    name = request.form.get("name", "").strip()
    phones = load_phones()
    data = load_data()

    # 🛑 If no data exists
    if not data:
        return render_template("not_found.html", name=name)

    last = data[-1]

    # 🔍 Exact match search (same logic, just safer)
    for i in range(len(last["names"])):
        if last["names"][i].strip().lower() == name.lower():
            return render_template(
                "person.html",
                name=last["names"][i],
                familycode=last["familycode"][i],
                address=last["addresses"][i],
                phone=phones.get(last["names"][i], "")
            )

    return render_template("not_found.html", name=name)

# -------------------------------
# CREATE (FIXED)
# -------------------------------

@app.route("/create", methods=["GET", "POST"])
def create():

    data = load_data()

    # 🧠 Always create a NEW list (do not delete old)
    amounts = [0] * len(namelist)

    new_entry = {
        "date": datetime.now().strftime("%d-%m-%Y %H:%M"),
        "names": namelist,
        "addresses": addresslist,
        "familycode": familycode,
        "amounts": amounts,
        "total": 0
    }

    data.append(new_entry)   # ✅ ADD new list (don’t replace)

    save_data(data)

    return redirect("/storage")

# -------------------------------
# STORAGE
# -------------------------------

@app.route("/storage")
def storage():
    return render_template("storage.html", data=load_data())

# -------------------------------
# DELETE (NEW FIX)
# -------------------------------

@app.route("/delete/<int:index>")
def delete(index):

    data = load_data()

    if 0 <= index < len(data):
        data.pop(index)
        save_data(data)

    return redirect("/storage")

# -------------------------------
# ANALYTICS
# -------------------------------

@app.route("/analytics")
def analytics():

    data = load_data()
    pending = []

    if data:
        last = data[-1]

        names = last.get("names", [])
        amounts = last.get("amounts", [])

        for i in range(min(len(names), len(amounts))):
            if amounts[i] == 0:
                pending.append(names[i])

    return render_template("analytics.html", pending=pending)

# -------------------------------
# DOWNLOAD (FIXED SAFE PATH)
# -------------------------------

@app.route("/download_prayer")
def download_prayer():

    doc = Document()
    doc.add_heading("Prayer Time - Darjeeling & Jalpaiguri", 0)

    table = doc.add_table(rows=1, cols=3)

    header = table.rows[0].cells
    header[0].text = "Month"
    header[1].text = "Morning"
    header[2].text = "Evening"

    data = [
        ("January","06:21 AM","05:05 PM"),
        ("February","06:08 AM","05:26 PM"),
        ("March","05:41 AM","05:43 PM"),
        ("April","05:10 AM","05:57 PM"),
        ("May","04:48 AM","06:13 PM"),
        ("June","04:42 AM","06:26 PM"),
        ("July","04:53 AM","06:26 PM"),
        ("August","05:07 AM","06:08 PM"),
        ("September","05:20 AM","05:37 PM"),
        ("October","05:34 AM","05:06 PM"),
        ("November","05:53 AM","04:45 PM"),
        ("December","06:13 AM","04:46 PM")
    ]

    for row_data in data:
        row = table.add_row().cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
        row[2].text = row_data[2]

    file_path = "prayer_table.docx"
    doc.save(file_path)

    return send_file(file_path, as_attachment=True)

# -------------------------------
# 📱 ADD PHONE (SAVE BUTTON FIX)
# -------------------------------

@app.route("/add_phone", methods=["POST"])
def add_phone():

    name = request.form.get("name")
    phone = request.form.get("phone", "").strip()

    # ✅ STRICT VALIDATION
    if not phone.isdigit() or len(phone) != 10:
        return "❌ Phone number must be exactly 10 digits"

    phones = load_phones()
    phones[name] = phone
    save_phones(phones)

    return redirect("/analytics")

# -------------------------------
# 🔍 SEARCH FOR ANALYTICS (AJAX FIX)
# -------------------------------

@app.route("/search_name", methods=["POST"])
def search_name():

    name = request.form.get("name", "").strip()
    data = load_data()
    phones = load_phones()

    if not data:
        return jsonify({"status": "not_found"})

    last = data[-1]

    for i in range(len(last["names"])):
        if last["names"][i].strip().lower() == name.lower():

            return jsonify({
                "status": "found",
                "name": last["names"][i],
                "familycode": last["familycode"][i],
                "address": last["addresses"][i],
                "phone": phones.get(last["names"][i], "")
            })

    return jsonify({"status": "not_found"})

# -------------------------------
# ⬇ DOWNLOAD LIST (FIXED)
# -------------------------------

# -------------------------------
# ⬇ DOWNLOAD DOCX (WITH YOUR FORMAT)
# -------------------------------
from docx.shared import Inches
import os

@app.route("/download/<int:index>")
def download(index):

    data = load_data()

    if index >= len(data):
        return "Invalid index"

    item = data[index]

    # 🔥 MAP YOUR EXISTING DATA (NO LOGIC CHANGE)
    namelist = item["names"]
    addresslist = item["addresses"]
    familycode = item["familycode"]
    amountlist = item["amounts"]
    total = item["total"]

    # -------------------------------
    # CREATE WORD FILE (YOUR CODE)
    # -------------------------------

    doc = Document()

    doc.add_heading("SATSANG PHILANTHROPY, SATSANG DEOGHAR", 1)
    doc.add_paragraph("JYOTINAGAR, PHANSIDEWA UPOYOJANA KENDRA-43")
    doc.add_paragraph("TOTAL AMOUNT - __________________________")
    doc.add_paragraph("(Rupees __________________________________ only)")
    doc.add_paragraph("POWER JYOTI CHALLAN, SBI, RANIDANGA")
    doc.add_paragraph("JOURNAL NO. - __________________     Dtd - __________________")
    doc.add_paragraph(f"Start Srl No - 01     End Srl No - {len(namelist)}     Deposited By: NITYANANDA SHIL")
    doc.add_paragraph("Family Code: 000005193400")

    doc.add_paragraph(" ")

    # -------------------------------
    # TABLE (YOUR CODE)
    # -------------------------------

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.autofit = False

    for row in table.rows:
        row.cells[0].width = Inches(0.5)
        row.cells[1].width = Inches(2.2)
        row.cells[2].width = Inches(2.5)
        row.cells[3].width = Inches(1.8)
        row.cells[4].width = Inches(1.2)

    header = table.rows[0].cells
    header[0].text = "Srl No"
    header[1].text = "Name"
    header[2].text = "Address"
    header[3].text = "Family Code"
    header[4].text = "Amount"

    previous_address = ""

    for i in range(len(namelist)):
        row = table.add_row().cells

        row[0].text = str(i+1)
        row[1].text = namelist[i]

        if addresslist[i] == previous_address:
            row[2].text = '"'
        else:
            row[2].text = addresslist[i]
            previous_address = addresslist[i]

        row[3].text = familycode[i]
        row[4].text = str(amountlist[i])

    # -------------------------------
    # TOTAL + FOOTER (YOUR CODE)
    # -------------------------------

    doc.add_paragraph("\nTOTAL AMOUNT = " + str(total))

    p = doc.add_paragraph("\nNIRMALLYA SHIL, JYOTINAGAR, PHANSIDEWA, DARJEELING, PIN-734434, PH. NO.- 8158002724")
    p.alignment = 1

    # -------------------------------
    # SAVE + DOWNLOAD
    # -------------------------------

    file_path = os.path.join(os.getcwd(), f"satsang_list_{index}.docx")
    doc.save(file_path)

    return send_file(file_path, as_attachment=True)

# -------------------------------
# 💰 DEPOSIT (FIX)
# -------------------------------

@app.route("/deposit", methods=["POST"])
def deposit():

    name = request.form.get("name")
    amount = request.form.get("amount")

    if not amount:
        return redirect("/analytics")

    amount = int(amount)

    data = load_data()

    if not data:
        return redirect("/analytics")

    last = data[-1]

    for i in range(len(last["names"])):
        if last["names"][i] == name:
            last["amounts"][i] += amount
            break

    # 🔥 Update total
    last["total"] = sum(last["amounts"])

    save_data(data)

    return redirect("/analytics")

@app.route("/duplicates")
def duplicates():

    from collections import Counter

    data = load_data()

    if not data:
        return render_template("duplicates.html", duplicates=[])

    last = data[-1]
    names = last["names"]

    count = Counter(names)

    duplicate_names = [name for name, c in count.items() if c > 1]

    return render_template("duplicates.html", duplicates=duplicate_names)

# -------------------------------
# 📱 PHONE PAGE
# -------------------------------

@app.route("/phone")
def phone_page():

    data = load_data()

    if not data:
        return render_template("phone.html", names=[])

    last = data[-1]

    return render_template("phone.html", names=last["names"])


# -------------------------------
# 📡 GET PHONE (AJAX)
# -------------------------------

@app.route("/get_phone", methods=["POST"])
def get_phone():

    name = request.form.get("name", "").strip()
    phones = load_phones()

    return jsonify({
        "name": name,
        "phone": phones.get(name, "")
    })


# -------------------------------
# 🔄 UPDATE PHONE
# -------------------------------

@app.route("/update_phone", methods=["POST"])
def update_phone():

    name = request.form.get("name", "").strip()
    phone = request.form.get("phone", "").strip()

    # ✅ STRICT VALIDATION (VERY IMPORTANT)
    if not phone.isdigit() or len(phone) != 10:
        return "❌ Phone must be exactly 10 digits"

    phones = load_phones()

    phones[name] = phone
    save_phones(phones)

    return "success"

# -------------------------------
# ✏️ CHANGE AMOUNT PAGE
# -------------------------------

@app.route("/change")
def change_page():

    data = load_data()

    if not data:
        return render_template("change.html", names=[])

    last = data[-1]

    return render_template("change.html", names=last["names"])


# -------------------------------
# 🔍 GET PERSON DATA (AJAX)
# -------------------------------

@app.route("/get_amount", methods=["POST"])
def get_amount():

    name = request.form.get("name")
    data = load_data()

    if not data:
        return jsonify({"status": "not_found"})

    last = data[-1]

    for i in range(len(last["names"])):
        if last["names"][i] == name:
            return jsonify({
                "status": "found",
                "name": name,
                "amount": last["amounts"][i]
            })

    return jsonify({"status": "not_found"})


# -------------------------------
# 🔁 UPDATE AMOUNT
# -------------------------------

@app.route("/update_amount", methods=["POST"])
def update_amount():

    name = request.form.get("name")
    new_amount = request.form.get("amount")

    if not new_amount.isdigit():
        return "❌ Enter valid number"

    new_amount = int(new_amount)

    data = load_data()

    if not data:
        return redirect("/change")

    last = data[-1]

    for i in range(len(last["names"])):

        if last["names"][i] == name:

            if last["amounts"][i] == 0:
                return "❌ Cannot change. No deposit found."

            last["amounts"][i] = new_amount
            break

    # update total
    last["total"] = sum(last["amounts"])

    save_data(data)

    return redirect("/change")
import os

@app.route("/debug")
def debug():
    return str(os.listdir("templates"))
@app.route("/save_month_status", methods=["POST"])
def save_month_status():
    month = request.form.get("month")
    status = request.form.get("status")

    try:
        with open("month_status.json", "r") as f:
            data = json.load(f)
    except:
        data = {}

    data[month] = status

    with open("month_status.json", "w") as f:
        json.dump(data, f, indent=4)

    return "success"

@app.route("/update_satsangi")
def update_satsangi():
    data = load_data()

    if not data:
        return render_template("update_satsangi.html", names=[])

    last = data[-1]
    return render_template("update_satsangi.html", names=last["names"])

@app.route("/get_satsangi", methods=["POST"])
def get_satsangi():
    name = request.form.get("name")
    data = load_data()

    if not data:
        return jsonify({"status": "not_found"})

    last = data[-1]

    for i in range(len(last["names"])):
        if last["names"][i] == name:
            return jsonify({
                "status": "found",
                "name": name,
                "address": last["addresses"][i],
                "familycode": last["familycode"][i]
            })

    return jsonify({"status": "not_found"})

@app.route("/delete_satsangi", methods=["POST"])
def delete_satsangi():
    name = request.form.get("name")
    data = load_data()

    if not data:
        return "error"

    last = data[-1]

    if name in last["names"]:
        i = last["names"].index(name)

        last["names"].pop(i)
        last["addresses"].pop(i)
        last["familycode"].pop(i)
        last["amounts"].pop(i)

        save_data(data)

    return "success"

@app.route("/add_satsangi", methods=["POST"])
def add_satsangi():
    name = request.form.get("name")
    address = request.form.get("address")
    familycode = request.form.get("familycode")

    data = load_data()

    if not data:
        return "error"

    last = data[-1]

    last["names"].append(name)
    last["addresses"].append(address)
    last["familycode"].append(familycode)
    last["amounts"].append(0)

    save_data(data)

    return "success"

# -------------------------------
# RUN
# -------------------------------

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)